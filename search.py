"""
File Search Utility - EXACT MATCH (Case Insensitive)
Searches for EXACT word matches inside Excel (.xlsx, .xls), Word (.docx), CSV files,
and Google Docs (.gdoc) / Google Sheets (.gsheet) files.
Example: searching for "Anna" will find "anna", "Anna" but NOT "Joanna" or "annabel"
"""

import os
import re
import warnings
from pathlib import Path
from typing import List, Optional, Set, Tuple
from dataclasses import dataclass, field
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor, as_completed

# Suppress openpyxl warnings about missing default styles
warnings.filterwarnings('ignore', category=UserWarning, module='openpyxl')

# ============================================================
# HARDCODE YOUR BASE PATH HERE
# ============================================================
# For Google Colab with Google Drive:
BASE_PATH = "/content/drive/Shared drives"

# For local Windows:
# BASE_PATH = "C:/Users/avitkin/Documents"

# Leave empty to use full paths:
# BASE_PATH = ""
# ============================================================

# Required libraries for reading Office files
try:
    import openpyxl
except ImportError:
    openpyxl = None

try:
    import xlrd
except ImportError:
    xlrd = None

try:
    from docx import Document
except ImportError:
    Document = None

# Google API support (for Google Docs and Sheets)
import json

try:
    from google.colab import auth
    from googleapiclient.discovery import build
    from googleapiclient.errors import HttpError
    GOOGLE_API_AVAILABLE = True
except ImportError:
    GOOGLE_API_AVAILABLE = False
    auth = None
    build = None
    HttpError = None

# Google API service objects (initialized lazily)
_sheets_service = None
_docs_service = None
_drive_service = None
_google_credentials = None
_google_authenticated = False


def init_google_services(silent: bool = False):
    """Initialize Google API services (authenticate once, reuse services).
    
    SECURITY: Uses READ-ONLY scopes only - cannot modify or delete any files.
    
    Args:
        silent: If True, don't print any messages
    """
    global _sheets_service, _docs_service, _drive_service, _google_credentials, _google_authenticated
    
    if not GOOGLE_API_AVAILABLE:
        return False
    
    if _google_authenticated:
        return True
    
    try:
        if not silent:
            print("\n  🔐 Google file detected - requesting READ-ONLY API access...")
        
        # Authenticate with Google using READ-ONLY scopes
        # These scopes CANNOT modify, delete, or create any files
        from google.colab import auth as colab_auth
        colab_auth.authenticate_user()
        
        # Import credentials with explicit READ-ONLY scopes
        import google.auth
        READONLY_SCOPES = [
            'https://www.googleapis.com/auth/drive.readonly',         # Read-only Drive access
            'https://www.googleapis.com/auth/spreadsheets.readonly',  # Read-only Sheets
            'https://www.googleapis.com/auth/documents.readonly',     # Read-only Docs
        ]
        
        credentials, project = google.auth.default(scopes=READONLY_SCOPES)
        _google_credentials = credentials
        
        # Build ALL service objects with read-only credentials
        _drive_service = build('drive', 'v3', credentials=credentials, cache_discovery=False)
        _sheets_service = build('sheets', 'v4', credentials=credentials, cache_discovery=False)
        _docs_service = build('docs', 'v1', credentials=credentials, cache_discovery=False)
        
        _google_authenticated = True
        if not silent:
            print("  ✓ Authenticated (READ-ONLY mode) - your files are safe\n")
        return True
    except Exception as e:
        if not silent:
            print(f"  ⚠️  Google API authentication failed: {e}\n")
        return False


def search_in_gsheet(file_path: Path, search_value: str) -> Optional[str]:
    """Search for EXACT value match inside a Google Sheet using Google Sheets API.
    
    SECURITY: Uses READ-ONLY API access.
    """
    global _sheets_service, _drive_service
    
    if not init_google_services():
        return None
    
    # Get file name (without extension) to search by title
    file_name = file_path.stem
    
    try:
        # Search for the file using READ-ONLY Drive API
        query = f"name = '{file_name}' and mimeType = 'application/vnd.google-apps.spreadsheet'"
        results = _drive_service.files().list(
            q=query,
            spaces='drive',
            fields='files(id, name)',
            supportsAllDrives=True,
            includeItemsFromAllDrives=True
        ).execute()
        
        files = results.get('files', [])
        
        if not files:
            return None
        
        file_id = files[0]['id']
        
        # Get spreadsheet data using READ-ONLY Sheets API
        spreadsheet = _sheets_service.spreadsheets().get(spreadsheetId=file_id).execute()
        sheets = spreadsheet.get('sheets', [])
        
        matches = []
        
        for sheet in sheets:
            sheet_name = sheet['properties']['title']
            
            # Read all values from the sheet
            range_name = f"'{sheet_name}'"
            result = _sheets_service.spreadsheets().values().get(
                spreadsheetId=file_id,
                range=range_name
            ).execute()
            
            values = result.get('values', [])
            for row_idx, row in enumerate(values, 1):
                for col_idx, cell_value in enumerate(row):
                    if cell_value:
                        if is_exact_match(str(cell_value), search_value):
                            # Convert to Excel-like column letter
                            col_letter = ""
                            col_num = col_idx + 1
                            while col_num > 0:
                                col_num, remainder = divmod(col_num - 1, 26)
                                col_letter = chr(65 + remainder) + col_letter
                            matches.append(f"Sheet '{sheet_name}', Cell {col_letter}{row_idx}")
        
        if matches:
            return "; ".join(matches[:5]) + (f" (+{len(matches)-5} more)" if len(matches) > 5 else "")
        return None
        
    except HttpError:
        return None
    except Exception:
        return None


def search_in_gdoc(file_path: Path, search_value: str) -> Optional[str]:
    """Search for EXACT value match inside a Google Doc using Google Docs API.
    
    SECURITY: Uses READ-ONLY API access.
    """
    global _docs_service, _drive_service
    
    if not init_google_services():
        return None
    
    # Get file name (without extension) to search by title
    file_name = file_path.stem
    
    try:
        # Search for the file using READ-ONLY Drive API
        query = f"name = '{file_name}' and mimeType = 'application/vnd.google-apps.document'"
        results = _drive_service.files().list(
            q=query,
            spaces='drive',
            fields='files(id, name)',
            supportsAllDrives=True,
            includeItemsFromAllDrives=True
        ).execute()
        
        files = results.get('files', [])
        
        if not files:
            return None
        
        file_id = files[0]['id']
        
        # Get document content using READ-ONLY Docs API
        document = _docs_service.documents().get(documentId=file_id).execute()
        
        matches = []
        
        # Extract text from body content
        body = document.get('body', {})
        content = body.get('content', [])
        
        paragraph_num = 0
        for element in content:
            if 'paragraph' in element:
                paragraph_num += 1
                paragraph = element['paragraph']
                elements = paragraph.get('elements', [])
                
                para_text = ""
                for elem in elements:
                    if 'textRun' in elem:
                        para_text += elem['textRun'].get('content', '')
                
                if para_text and is_exact_match(para_text, search_value):
                    preview = para_text[:50].replace('\n', ' ') + "..." if len(para_text) > 50 else para_text.replace('\n', ' ')
                    matches.append(f"Paragraph {paragraph_num}: '{preview}'")
            
            elif 'table' in element:
                table = element['table']
                table_rows = table.get('tableRows', [])
                
                for row_idx, row in enumerate(table_rows):
                    cells = row.get('tableCells', [])
                    for cell_idx, cell in enumerate(cells):
                        cell_content = cell.get('content', [])
                        cell_text = ""
                        for cell_elem in cell_content:
                            if 'paragraph' in cell_elem:
                                for elem in cell_elem['paragraph'].get('elements', []):
                                    if 'textRun' in elem:
                                        cell_text += elem['textRun'].get('content', '')
                        
                        if cell_text and is_exact_match(cell_text, search_value):
                            matches.append(f"Table Row {row_idx + 1}, Cell {cell_idx + 1}")
        
        if matches:
            return "; ".join(matches[:5]) + (f" (+{len(matches)-5} more)" if len(matches) > 5 else "")
        return None
        
    except HttpError:
        return None
    except Exception:
        return None


# Supported file extensions
SUPPORTED_EXTENSIONS = {'.xlsx', '.xls', '.docx', '.csv', '.gsheet', '.gdoc'}


def convert_windows_path(path: str) -> str:
    """
    Convert Windows path to Google Drive compatible path.

    """
    # Remove common Windows drive prefixes
    prefixes_to_remove = [
        r'G:\Shared drives\\',
        r'G:\Shared drives/',
        r'G:/Shared drives/',
        r'G:/Shared drives\\',
        'G:\\Shared drives\\',
        'G:/Shared drives/',
    ]
    
    result = path
    for prefix in prefixes_to_remove:
        if result.startswith(prefix):
            result = result[len(prefix):]
            break
    
    # Replace backslashes with forward slashes
    result = result.replace('\\', '/')
    
    # Remove leading/trailing slashes
    result = result.strip('/')
    
    return result


@dataclass
class FolderResult:
    """Represents search results for a final folder."""
    folder_name: str
    folder_path: str
    all_files: List[str]
    searched_file: Optional[str]
    searched_file_modified: Optional[str]
    search_found: bool
    search_details: Optional[str]


def check_dependencies():
    """Check if required libraries are installed."""
    missing = []
    if openpyxl is None:
        missing.append("openpyxl")
    if xlrd is None:
        missing.append("xlrd")
    if Document is None:
        missing.append("python-docx")
    
    if missing:
        print("Missing required libraries. Install them with:")
        print(f"  pip install {' '.join(missing)}")
        return False
    
    return True


# Cache for compiled regex patterns (thread-safe for reading)
_pattern_cache: dict = {}


def get_compiled_pattern(search_value: str) -> re.Pattern:
    """Get or create a compiled regex pattern for the search value."""
    if search_value not in _pattern_cache:
        # Use word boundary regex for exact matching (case insensitive)
        # \b matches word boundaries (start/end of word)
        pattern = r'\b' + re.escape(search_value) + r'\b'
        _pattern_cache[search_value] = re.compile(pattern, re.IGNORECASE)
    return _pattern_cache[search_value]


def is_exact_match(cell_value: str, search_value: str) -> bool:
    """
    Check if search_value is an EXACT match in cell_value (case insensitive).
    Uses compiled regex for better performance.
    
    Examples:
        is_exact_match("Anna", "anna") -> True
        is_exact_match("anna", "Anna") -> True  
        is_exact_match("Joanna", "anna") -> False
        is_exact_match("annabel", "anna") -> False
        is_exact_match("Hello Anna!", "anna") -> True
        is_exact_match("Anna,Bob", "anna") -> True
    """
    compiled_pattern = get_compiled_pattern(search_value)
    return bool(compiled_pattern.search(cell_value))


def search_in_xlsx(file_path: Path, search_value: str) -> Optional[str]:
    """Search for EXACT value match inside an .xlsx file."""
    if openpyxl is None:
        return None
    
    try:
        workbook = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        matches = []
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value is not None:
                        cell_str = str(cell.value)
                        if is_exact_match(cell_str, search_value):
                            matches.append(f"Sheet '{sheet_name}', Cell {cell.coordinate}")
        
        workbook.close()
        
        if matches:
            return "; ".join(matches[:5]) + (f" (+{len(matches)-5} more)" if len(matches) > 5 else "")
        return None
        
    except Exception:
        return None


def search_in_xls(file_path: Path, search_value: str) -> Optional[str]:
    """Search for EXACT value match inside an .xls file (older Excel format)."""
    if xlrd is None:
        return None
    
    try:
        workbook = xlrd.open_workbook(str(file_path))
        matches = []
        
        for sheet_idx in range(workbook.nsheets):
            sheet = workbook.sheet_by_index(sheet_idx)
            sheet_name = sheet.name
            
            for row_idx in range(sheet.nrows):
                for col_idx in range(sheet.ncols):
                    cell_value = sheet.cell_value(row_idx, col_idx)
                    if cell_value:
                        cell_str = str(cell_value)
                        if is_exact_match(cell_str, search_value):
                            col_letter = xlrd.colname(col_idx)
                            matches.append(f"Sheet '{sheet_name}', Cell {col_letter}{row_idx + 1}")
        
        if matches:
            return "; ".join(matches[:5]) + (f" (+{len(matches)-5} more)" if len(matches) > 5 else "")
        return None
        
    except Exception:
        return None


def search_in_docx(file_path: Path, search_value: str) -> Optional[str]:
    """Search for EXACT value match inside a .docx file."""
    if Document is None:
        return None
    
    try:
        doc = Document(str(file_path))
        matches = []
        
        # Search in paragraphs
        for i, para in enumerate(doc.paragraphs, 1):
            text = para.text
            if is_exact_match(text, search_value):
                preview = text[:50] + "..." if len(text) > 50 else text
                matches.append(f"Paragraph {i}: '{preview}'")
        
        # Search in tables
        for table_idx, table in enumerate(doc.tables, 1):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    text = cell.text
                    if is_exact_match(text, search_value):
                        matches.append(f"Table {table_idx}, Row {row_idx + 1}, Cell {cell_idx + 1}")
        
        if matches:
            return "; ".join(matches[:5]) + (f" (+{len(matches)-5} more)" if len(matches) > 5 else "")
        return None
        
    except Exception:
        return None


def search_in_csv(file_path: Path, search_value: str) -> Optional[str]:
    """Search for EXACT value match inside a .csv file."""
    
    try:
        # Read file as binary first, then decode
        file_path_str = str(file_path)
        
        # Read raw bytes
        with open(file_path_str, 'rb') as f:
            raw_data = f.read()
        
        # Try to decode with different encodings
        content = None
        for encoding in ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'iso-8859-1']:
            try:
                content = raw_data.decode(encoding)
                break
            except:
                continue
        
        if not content:
            # Last resort: decode with errors ignored
            content = raw_data.decode('utf-8', errors='ignore')
        
        if not content:
            return None
        
        matches = []
        
        # Search line by line (simpler and more reliable)
        lines = content.replace('\r\n', '\n').replace('\r', '\n').split('\n')
        
        for row_idx, line in enumerate(lines, 1):
            if not line.strip():
                continue
            
            # First check if the line might contain a match
            if is_exact_match(line, search_value):
                # Try to find which column
                cells = line.split(',')
                found_in_cell = False
                for col_idx, cell in enumerate(cells):
                    cell_clean = cell.strip().strip('"').strip("'")
                    if is_exact_match(cell_clean, search_value):
                        # Convert to Excel-like column letter
                        col_letter = ""
                        col_num = col_idx + 1
                        while col_num > 0:
                            col_num, remainder = divmod(col_num - 1, 26)
                            col_letter = chr(65 + remainder) + col_letter
                        matches.append(f"Row {row_idx}, Col {col_letter}")
                        found_in_cell = True
                        break
                
                if not found_in_cell:
                    # If we can't find specific column, just report the row
                    matches.append(f"Row {row_idx}")
        
        if matches:
            return "; ".join(matches[:5]) + (f" (+{len(matches)-5} more)" if len(matches) > 5 else "")
        return None
        
    except Exception as e:
        # Uncomment next line for debugging:
        # print(f"  [DEBUG] CSV error for {file_path}: {e}")
        return None


def search_in_file(file_path: Path, search_value: str) -> Optional[str]:
    """Search for EXACT value match in a file based on its extension."""
    extension = file_path.suffix.lower()
    
    if extension == '.xlsx':
        return search_in_xlsx(file_path, search_value)
    elif extension == '.xls':
        return search_in_xls(file_path, search_value)
    elif extension == '.docx':
        return search_in_docx(file_path, search_value)
    elif extension == '.csv':
        return search_in_csv(file_path, search_value)
    elif extension == '.gsheet':
        return search_in_gsheet(file_path, search_value)
    elif extension == '.gdoc':
        return search_in_gdoc(file_path, search_value)
    
    return None


# Folders to skip during search (case-insensitive)
SKIP_FOLDERS = {'old', 'test', 'bug', 'feasibility'}


def is_skip_folder(folder: Path) -> bool:
    """Check if folder should be skipped (case-insensitive)."""
    return folder.name.lower() in SKIP_FOLDERS


def is_old_folder(folder: Path) -> bool:
    """Check if folder is named 'Old' (case-insensitive). Kept for backwards compatibility."""
    return folder.name.lower() == 'old'


def has_non_skip_subfolders(folder: Path) -> bool:
    """Check if folder has any subfolders (excluding skip folders like 'Old', 'test', etc.)."""
    try:
        for item in folder.iterdir():
            if item.is_dir() and not is_skip_folder(item):
                return True
    except PermissionError:
        pass
    return False


def is_final_folder(folder: Path) -> bool:
    """Check if folder is a 'final' folder (has no subfolders except possibly skip folders)."""
    return not has_non_skip_subfolders(folder)


def get_supported_files(folder: Path) -> List[Path]:
    """Get all supported files in a folder (not recursive)."""
    files = []
    try:
        for item in folder.iterdir():
            if item.is_file() and item.suffix.lower() in SUPPORTED_EXTENSIONS:
                files.append(item)
    except PermissionError:
        pass
    return files


def get_most_recent_file(files: List[Path]) -> Optional[Path]:
    """Get the most recently modified file from a list."""
    if not files:
        return None
    return max(files, key=lambda f: f.stat().st_mtime)


def scan_subtree(root: Path) -> Set[Path]:
    """Scan a subtree for final folders - used for parallel execution."""
    final_folders: Set[Path] = set()
    
    if not root.exists() or not root.is_dir():
        return final_folders
    
    # Use os.walk for traversal with directory skipping
    for dirpath, dirnames, filenames in os.walk(root):
        current = Path(dirpath)
        
        # Remove skip folders from dirnames to prevent descending into them
        dirnames[:] = [d for d in dirnames if d.lower() not in SKIP_FOLDERS]
        
        # Check if current folder has no non-skip subdirectories (is final)
        if not dirnames:
            final_folders.add(current)
    
    return final_folders


def get_top_level_subdirs(root: Path) -> List[Path]:
    """Get immediate subdirectories of a path (excluding skip folders)."""
    subdirs = []
    try:
        for item in root.iterdir():
            if item.is_dir() and item.name.lower() not in SKIP_FOLDERS:
                subdirs.append(item)
    except PermissionError:
        pass
    return subdirs


def find_all_final_folders(root_paths: List[str]) -> List[Path]:
    """Find all final folders using parallel scanning."""
    all_final_folders: Set[Path] = set()
    valid_roots = []
    
    # Validate paths first
    for root_path in root_paths:
        root_path = root_path.strip()
        if not root_path:
            continue
        root = Path(root_path)
        if not root.exists():
            print(f"  ⚠️  Warning: Folder not found: {root_path}")
            continue
        if not root.is_dir():
            print(f"  ⚠️  Warning: Path is not a directory: {root_path}")
            continue
        valid_roots.append(root)
        print(f"  📂 Queued for scanning: {root_path}")
    
    if not valid_roots:
        return []
    
    # Collect all subtrees to scan in parallel
    # This gives parallelism even with a single root path
    subtrees_to_scan: List[Path] = []
    
    for root in valid_roots:
        top_subdirs = get_top_level_subdirs(root)
        if top_subdirs:
            # Has subdirectories - add them for parallel scanning
            subtrees_to_scan.extend(top_subdirs)
        else:
            # Root itself is a final folder
            all_final_folders.add(root)
    
    if not subtrees_to_scan:
        return sorted(all_final_folders, key=lambda f: str(f).lower())
    
    print(f"\n  ⚡ Scanning {len(subtrees_to_scan)} subtrees in parallel...")
    
    # Scan all subtrees in parallel
    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
        futures = {executor.submit(scan_subtree, subtree): subtree for subtree in subtrees_to_scan}
        
        completed = 0
        total = len(futures)
        for future in as_completed(futures):
            completed += 1
            subtree = futures[future]
            try:
                folders = future.result()
                all_final_folders.update(folders)
                # Show progress
                if total > 10:
                    progress = int((completed / total) * 20)
                    bar = "█" * progress + "░" * (20 - progress)
                    print(f"  ⏳ [{bar}] {completed}/{total} subtrees scanned", end='\r')
            except Exception:
                pass
    
    if total > 10:
        print()  # New line after progress bar
    
    print(f"  ✓ Found {len(all_final_folders)} final folders")
    
    return sorted(all_final_folders, key=lambda f: str(f).lower())


# Number of parallel workers for file searching
# - Local SSD: 4-8 workers
# - Local HDD: 4-6 workers  
# - Network/Google Drive: 16-32 workers (more I/O latency = more parallelism helps)
MAX_WORKERS = 16


def process_single_folder(folder: Path, search_value: str) -> Optional[FolderResult]:
    """Process a single folder - used for parallel execution."""
    files = get_supported_files(folder)
    all_file_names = sorted([f.name for f in files])
    
    # Skip empty folders (no supported files)
    if not files:
        return None
    
    most_recent = get_most_recent_file(files)
    
    if most_recent:
        mod_time = datetime.fromtimestamp(most_recent.stat().st_mtime)
        mod_time_str = mod_time.strftime("%Y-%m-%d %H:%M:%S")
        
        search_result = search_in_file(most_recent, search_value)
        
        return FolderResult(
            folder_name=folder.name,
            folder_path=str(folder),
            all_files=all_file_names,
            searched_file=most_recent.name,
            searched_file_modified=mod_time_str,
            search_found=search_result is not None,
            search_details=search_result
        )
    else:
        return FolderResult(
            folder_name=folder.name,
            folder_path=str(folder),
            all_files=all_file_names,
            searched_file=None,
            searched_file_modified=None,
            search_found=False,
            search_details=None
        )


def search_in_final_folders(
    search_value: str,
    folder_paths: List[str]
) -> List[FolderResult]:
    """Search for EXACT value match in final folders."""
    
    results: List[FolderResult] = []
    final_folders = find_all_final_folders(folder_paths)
    
    if not final_folders:
        print("\n  ⚠️  No final folders found.")
        return results
    
    total = len(final_folders)
    print(f"\n  🔍 Searching {total} folder(s)...")
    
    # Process folders with progress bar
    for idx, folder in enumerate(final_folders, 1):
        try:
            # Show progress bar
            progress = int((idx / total) * 30)
            bar = "█" * progress + "░" * (30 - progress)
            print(f"\r  ⏳ [{bar}] {idx}/{total}", end="", flush=True)
            
            result = process_single_folder(folder, search_value)
            if result is not None:
                results.append(result)
        except Exception:
            pass  # Skip folders with errors silently
    
    print(f"\r  ✓ Searched {total} folder(s)" + " " * 20)
    return results


def print_results(results: List[FolderResult], search_value: str) -> None:
    """Print search results in a formatted way."""
    if not results:
        print("\n❌ No folders processed.")
        return
    
    folders_with_matches = [r for r in results if r.search_found]
    
    print("\n")
    print("╔" + "═"*78 + "")
    print("║" + f"  🎯 EXACT MATCH RESULTS FOR: '{search_value}'".ljust(78) + "")
    print("╠" + "═"*78 + "")
    print("║" + f"  📁 Processed: {len(results)} final folder(s)".ljust(78) + "")
    print("║" + f"  ✅ Exact matches found in: {len(folders_with_matches)} folder(s)".ljust(78) + "")
    print("╚" + "═"*78 + "")
    
    if not folders_with_matches:
        print("\n  ⚠️  No exact matches found in any folder.")
        return
    
    print("\n")
    print("┌" + "─"*78 + "")
    print("│" + "  📂 FOLDERS WITH EXACT MATCHES".ljust(78) + "")
    print("└" + "─"*78 + "")
    
    for i, result in enumerate(folders_with_matches, 1):
        print(f"\n  ╭{'─'*74}")
        print(f"  │  {i}. 📁 {result.folder_name}".ljust(77) + "")
        print(f"  ├{'─'*74}")
        print(f"  │  📍 Path: {result.folder_path}".ljust(77) + "")
        print(f"  │".ljust(77) + "")
        print(f"  │  📋 Files in folder ({len(result.all_files)}):".ljust(77) + "")
        
        for fname in result.all_files:
            if fname == result.searched_file:
                line = f"  │      ★ {fname} ← SEARCHED (most recent)"
            else:
                line = f"  │      • {fname}"
            print(line.ljust(77) + "")
        
        print(f"  │".ljust(77))
        print(f"  ├{'─'*74}")
        
        details = result.search_details if result.search_details else ""
        print(f"  │  📄 Details: {details[:60]}{'...' if len(details) > 60 else ''}")
        print(f"  │".ljust(77))
        print(f"  ├{'─'*74}")
        
        full_file_path = f"{result.folder_path}/{result.searched_file}"
        print(f"  │  🎯 MATCH: {full_file_path}")
        print(f"  ╰{'─'*74}")


# Main function
if __name__ == "__main__":
    print("\n")
    print("  ╔═══════════════════════════════════════════════════════════════════")
    print("  ║                                                                   ")
    print("  ║   🎯  SEARCH UTILITY                                             ")
    print("  ║                                                                   ")
    print("  ║   📄 Searches in Excel, Word, CSV, Google Docs & Sheets          ")
    print("  ║   📂 Shows all files, searches the most recent one                ")
    print("  ║   🚫 Skips subfolders: Old, test, bug, feasibility                 ")
    print("  ║   ⚡ Parallel processing for faster search                        ")
    print("  ║                                                                   ")
    print("  ╚═══════════════════════════════════════════════════════════════════")
    print()
    
    if BASE_PATH:
        print(f"  📍 Base path: {BASE_PATH}")
        print("     (Enter folder names relative to this path)")
        print()
    
    if not check_dependencies():
        exit(1)
    
    search_value = ""
    while not search_value:
        search_value = input("  🎯 Enter the value to search for: ").strip()
        if not search_value:
            print("  ⚠️  Search value is required. Please try again.\n")
    
    folder_paths = []
    while not folder_paths:
        if BASE_PATH:
            print(f"\n  📁 Enter folder names (relative to base path)")
            print("     Press Enter on empty line when done:\n")
        else:
            print("\n  📁 Enter full folder paths to search")
            print("     Press Enter on empty line when done:\n")
        path_num = 1
        
        while True:
            folder = input(f"     [{path_num}] ➜ ").strip()
            if not folder:
                break
            
            # Convert Windows path if needed (e.g., G:\Shared drives\...)
            if folder.startswith("G:") or "\\" in folder:
                folder = convert_windows_path(folder)
                print(f"          → Converted to: {folder}")
            
            # Combine with BASE_PATH if set
            if BASE_PATH and not folder.startswith("/"):
                full_path = f"{BASE_PATH}/{folder}"
            else:
                full_path = folder
            folder_paths.append(full_path)
            path_num += 1
        
        if not folder_paths:
            print("  ⚠️  At least one folder path is required. Please try again.")
    
    print(f"\n  🚀 Searching for EXACT match of '{search_value}' in {len(folder_paths)} folder(s):")
    for fp in folder_paths:
        print(f"     • {fp}")
    print()
    
    try:
        results = search_in_final_folders(
            search_value=search_value,
            folder_paths=folder_paths
        )
        
        print_results(results, search_value)
        
        print("\n  ✅ Search completed!")
        print()
        
    except Exception as e:
        print(f"\n  ❌ An error occurred: {e}")

